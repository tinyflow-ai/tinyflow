from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

fonts = [
    "Arial Unicode MS",
    "Arial Unicode",
    "STHeiti",
    "Heiti SC",
    "Hiragino Sans GB",
    "Songti SC",
    "PingFang SC",
    "Microsoft YaHei",
    "OPPOSans",
    "MiSans",
    "SimSun",
]

doc = Document()
for font in fonts:
    p = doc.add_paragraph()
    r = p.add_run(f"{font}: 中文字体测试 - AIFlowy 智能体配置")
    r.font.name = font
    r.font.size = Pt(14)
    r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
doc.save("/Users/michael/git/tinyflow-ai-next/docs/aiflowy-guide/font-probe.docx")
