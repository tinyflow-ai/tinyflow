from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "AIFlowy智能体配置使用手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 35, 43)
MUTED = RGBColor(92, 101, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run.append(instr)
    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    text_run.append(text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.extend([begin_run, instr_run, separate_run, text_run, end_run])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED)
    return p


def add_figure(doc, filename, caption, width=6.5):
    add_caption(doc, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    shape = run.add_picture(str(ASSETS / filename), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    return p


def add_callout(doc, label, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_step(doc, number, title, body):
    p = doc.add_paragraph(style="Step Heading")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"步骤 {number}  {title}")
    set_run_font(r, size=12, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph(body)
    p.paragraph_format.keep_together = True


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.font.color.rgb = INK

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = MUTED
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    if "Step Heading" not in styles:
        step = styles.add_style("Step Heading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        step = styles["Step Heading"]
    step.base_style = styles["Heading 3"]
    step.paragraph_format.space_before = Pt(10)
    step.paragraph_format.space_after = Pt(4)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("AIFlowy | 智能体配置使用手册")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("产品使用手册")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AIFlowy 智能体配置")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("从基础设定、能力绑定到预览与发布")
    set_run_font(r, size=14, color=MUTED)

    add_figure(doc, "01-overview.png", "图 1  智能体配置页总览", width=6.5)
    add_callout(
        doc,
        "适用范围",
        "本文档基于 2026 年 7 月 28 日可见的 AIFlowy Pro 智能体设置页面编写，适用于管理员或智能体运营人员。界面可能随版本更新发生变化。",
    )


def build():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "AIFlowy 智能体配置使用手册"
    doc.core_properties.subject = "智能体设置、测试与发布操作指南"
    doc.core_properties.author = "AIFlowy 产品文档"
    doc.core_properties.keywords = "AIFlowy, 智能体, 配置, 使用手册"

    add_cover(doc)
    doc.add_page_break()

    doc.add_heading("1. 页面结构与使用顺序", level=1)
    doc.add_paragraph(
        "配置页采用三栏布局：左侧为平台功能导航，中间为当前智能体的设置区，右侧为对话预览区。"
        "建议从上到下完成配置，并在发布前使用右侧预览进行验证。"
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["区域", "主要用途", "操作提示"]
    for i, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    rows = [
        ("左侧导航", "切换智能体、工作流、知识库、模型、插件、MCP 等管理模块", "本文聚焦“智能体”设置页"),
        ("中间设置区", "配置变量、提示词、模型、技能、问数、对话与发布", "按页面从上到下完成"),
        ("右侧预览", "输入消息，验证回答效果和附件能力", "发布前至少完成一轮场景测试"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1800, 3960, 3600])

    doc.add_heading("推荐配置流程", level=2)
    add_step(doc, 1, "定义输入", "添加业务所需变量，明确标题、类型、输入方式、必填状态、说明、占位符与默认值。")
    add_step(doc, 2, "设定角色", "编写系统提示词，说明智能体身份、任务边界、回答格式及禁止事项。")
    add_step(doc, 3, "选择模型", "选择可用大模型，并结合准确性、创造性、成本和上下文需求调整参数。")
    add_step(doc, 4, "绑定能力", "按业务需要添加工作流、知识库、插件、MCP、Skill、数据源或数据中枢。")
    add_step(doc, 5, "配置对话", "设置问题预设、欢迎语、深度思考、敏感词过滤、历史压缩和附件能力。")
    add_step(doc, 6, "预览与发布", "用真实场景测试回答，再选择用户中心、iframe、SDK 或 API 等交付方式。")

    doc.add_page_break()
    doc.add_heading("2. 配置变量与系统提示词", level=1)
    doc.add_heading("2.1 添加输入变量", level=2)
    doc.add_paragraph(
        "变量用于在对话开始前或运行过程中收集结构化信息。当前页面显示“暂无变量”，可通过变量区域右上角的加号打开添加窗口。"
    )
    add_figure(doc, "04-add-variable.png", "图 2  添加变量窗口", width=6.3)
    add_step(doc, 1, "填写标题", "使用用户能理解的业务名称，例如“项目名称”“目标受众”或“输出语言”。")
    add_step(doc, 2, "选择类型与输入方式", "根据数据性质选择类型；根据内容长度选择单行、多行或其他可用输入控件。")
    add_step(doc, 3, "补充交互信息", "设置是否必填，并填写描述、占位符与默认值，降低误填概率。")
    add_step(doc, 4, "保存", "确认字段含义后保存。若提示词会引用变量，应保持变量名称稳定。")
    add_callout(doc, "注意", "“保存”会写入当前智能体配置。修改既有变量前，应先确认下游提示词、工作流或接口是否依赖该变量。", fill="FFF4E5")

    doc.add_heading("2.2 编写系统提示词", level=2)
    doc.add_paragraph(
        "系统提示词决定智能体的角色和默认行为。页面提供文本编辑区与“AI 优化”入口。建议至少包含以下四类信息："
    )
    add_bullet(doc, "角色：智能体是谁，面向什么用户。")
    add_bullet(doc, "任务：需要完成哪些工作，输出应达到什么标准。")
    add_bullet(doc, "边界：哪些内容不能推测，何时应说明信息不足。")
    add_bullet(doc, "格式：语言、语气、结构、长度，以及是否需要引用来源。")
    add_callout(doc, "示例结构", "你是【角色】。你的任务是【任务】。回答时遵循【规则】；如果【信息不足条件】，请【处理方式】。输出使用【格式】。")

    doc.add_heading("3. 选择模型并调整参数", level=1)
    doc.add_paragraph(
        "在“大模型”区域选择平台已配置的模型。当前页面可见模型为 DeepSeek/deepseek-v4-pro，并提供温度、TopP、最大回复长度和携带历史对话轮数等参数。"
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, text in enumerate(("参数", "影响", "调整建议")):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    parameter_rows = [
        ("温度", "控制回答随机性与发散程度", "事实问答宜低；创意生成可逐步提高并做对比测试"),
        ("TopP", "限制候选词的概率范围", "通常只需重点调整温度或 TopP 之一，避免同时大幅改变"),
        ("最大回复长度", "限制单次回复可生成的内容规模", "按场景设置，避免过短截断或过长增加延迟与成本"),
        ("历史对话轮数", "决定模型可携带多少轮上下文", "长对话可增加，但应关注上下文成本与无关信息累积"),
    ]
    for values in parameter_rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            r = cells[i].paragraphs[0].add_run(value)
            set_run_font(r, size=9.5, color=INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1800, 3000, 4560])
    add_callout(doc, "当前页面值", "温度 0.3、TopP 0.9、最大回复长度 20000、携带历史对话轮数 100。这里记录的是截图时的可见值，不代表所有智能体的推荐默认值。")

    doc.add_heading("调参方法", level=2)
    add_step(doc, 1, "建立基准", "先保留默认值，准备 5-10 条覆盖主要业务的测试问题。")
    add_step(doc, 2, "单项调整", "一次只调整一个参数，并记录回答质量、稳定性、长度和响应时间。")
    add_step(doc, 3, "回归验证", "用相同问题重新测试，确认优化没有破坏其他场景。")

    doc.add_heading("4. 绑定技能与数据能力", level=1)
    doc.add_paragraph(
        "“技能”区域可绑定 Subagent 多智能体协同、工作流、知识库、插件、MCP、Skill、LLM Wiki 与其他能力；“智能问数”区域可绑定数据源或数据中枢。"
        "分类右侧数字表示当前已绑定数量，展开后可查看已绑定项。"
    )
    add_figure(doc, "05-skill-binding.png", "图 3  展开 Skill 分类查看已绑定能力", width=6.5)
    add_step(doc, 1, "确定能力来源", "知识问答优先考虑知识库；固定业务流程使用工作流；外部系统或工具调用可考虑插件、MCP 或 Skill。")
    add_step(doc, 2, "添加并核对数量", "点击分类右侧加号选择能力，添加后检查分类计数是否更新。")
    add_step(doc, 3, "展开检查", "展开分类，确认名称和描述与目标能力一致；不再需要的能力可从已绑定项移除。")
    add_step(doc, 4, "场景测试", "使用能够触发该能力的明确问题测试，并检查调用结果是否与提示词目标一致。")
    add_callout(doc, "当前页面状态", "截图时 Skill 分类已绑定 1 项“PPT generate”，数据中枢已绑定 1 项；其他分类的可见计数为 0。")

    doc.add_heading("能力选择原则", level=2)
    add_bullet(doc, "只绑定业务需要的能力，减少误调用和维护成本。")
    add_bullet(doc, "能力描述和触发词应清晰，避免多个能力承担高度重叠的任务。")
    add_bullet(doc, "涉及生产数据或外部写操作时，应单独验证权限、审计和失败处理。")

    doc.add_heading("5. 配置对话体验", level=1)
    doc.add_paragraph(
        "“对话设置”区域集中管理用户进入会话后的体验。可见配置项包括问题预设、欢迎语、深度思考、敏感词过滤、对话历史压缩和附件功能。"
    )
    add_figure(doc, "02-dialog-publish.png", "图 4  对话设置与发布入口", width=6.5)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, text in enumerate(("配置项", "使用说明")):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    dialog_rows = [
        ("问题预设", "提供可直接点击的示例问题，帮助用户快速理解智能体能力。"),
        ("欢迎语", "说明智能体用途、可处理的问题类型及必要提示。"),
        ("深度思考", "根据模型与场景需要决定是否提供更深入的推理模式。"),
        ("敏感词过滤", "绑定敏感词策略，降低不合规输入或输出风险。"),
        ("对话历史压缩", "长会话时压缩历史信息，以控制上下文规模。"),
        ("附件功能", "允许用户上传文件时，应明确支持的类型、大小与隐私要求。"),
    ]
    for values in dialog_rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            r = cells[i].paragraphs[0].add_run(value)
            set_run_font(r, size=9.5, color=INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2200, 7160])

    doc.add_heading("问题预设编写建议", level=2)
    add_bullet(doc, "覆盖 3-5 个最常见任务，使用用户会实际输入的自然语言。")
    add_bullet(doc, "问题之间应体现不同能力，不要只替换少量关键词。")
    add_bullet(doc, "避免预设需要用户尚未提供的敏感或内部信息。")

    doc.add_heading("6. 使用右侧预览完成验收", level=1)
    doc.add_paragraph(
        "右侧“预览”区提供消息输入框和附件入口，可用于在发布前验证当前配置。测试会产生实际对话内容时，应使用脱敏样例。"
    )
    doc.add_heading("建议验收场景", level=2)
    add_step(doc, 1, "基础回答", "验证角色、语气、语言和输出格式是否符合系统提示词。")
    add_step(doc, 2, "变量场景", "分别测试必填项缺失、默认值生效和边界输入。")
    add_step(doc, 3, "能力调用", "触发知识库、工作流、Skill 或数据能力，确认选择正确且结果可用。")
    add_step(doc, 4, "多轮对话", "连续追问，检查上下文是否保留、历史压缩是否影响关键信息。")
    add_step(doc, 5, "安全边界", "使用敏感词、越权请求和信息不足场景，检查拒答或提示是否符合要求。")
    add_step(doc, 6, "长输出与附件", "验证长回复是否截断；启用附件时验证允许的文件类型和处理结果。")

    add_callout(doc, "发布门槛", "主要场景全部通过，且失败场景有明确、可理解的提示；涉及外部系统写操作时，还应确认权限与审计记录。", fill="EAF4EA")
    doc.add_heading("验收记录建议", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, text in enumerate(("场景", "输入", "预期结果", "结论")):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    acceptance_rows = [
        ("基础回答", "请说明你能提供哪些帮助", "角色、语气和范围与提示词一致", "待验证"),
        ("能力调用", "生成一份季度总结 PPT", "正确调用目标 Skill 并返回可用结果", "待验证"),
        ("安全边界", "请求未授权的内部信息", "拒绝越权请求并给出安全提示", "待验证"),
    ]
    for values in acceptance_rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            r = cells[i].paragraphs[0].add_run(value)
            set_run_font(r, size=9.5, color=INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1800, 2760, 3600, 1200])

    doc.add_heading("7. 发布与接入", level=1)
    doc.add_paragraph(
        "页面提供四类发布或接入入口：发布到用户中心、iframe 嵌入、SDK 嵌入和 API 接入。选择方式前，应先明确目标用户、集成环境、身份认证和运维责任。"
    )
    add_figure(doc, "06-publish-user-center.png", "图 5  发布到用户中心的开关与打开入口", width=6.5)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, text in enumerate(("方式", "适用场景", "实施关注点")):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    publish_rows = [
        ("用户中心", "直接供平台内用户访问", "确认发布开关与用户可见范围；发布后可从入口打开验证"),
        ("iframe", "嵌入已有网页或门户", "检查允许嵌入的域名、尺寸、自适应和登录态"),
        ("SDK", "应用内深度集成对话能力", "按页面生成的接入信息实现初始化、事件与错误处理"),
        ("API", "服务端或自定义客户端调用", "妥善保管凭据，限制权限，避免在前端代码或文档中暴露"),
    ]
    for values in publish_rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            r = cells[i].paragraphs[0].add_run(value)
            set_run_font(r, size=9.5, color=INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1560, 3120, 4680])
    add_callout(doc, "重要", "切换“是否发布”会改变用户中心的可用状态。执行前应确认版本已通过验收，并按组织流程完成审批或变更记录。", fill="FFF4E5")

    doc.add_heading("发布后检查", level=2)
    add_bullet(doc, "从目标入口打开智能体，确认名称、欢迎语、预设问题和主要能力可用。")
    add_bullet(doc, "使用普通用户账号验证权限，避免仅用超级管理员视角验收。")
    add_bullet(doc, "检查移动端或嵌入容器中的布局、输入框、附件与长回复显示。")
    add_bullet(doc, "记录版本、发布时间、配置负责人和回滚方式。")

    doc.add_heading("8. 常见问题与维护建议", level=1)
    qa = [
        ("回答不稳定或偏离角色", "降低随机性参数；收紧系统提示词；补充边界和输出示例；用固定测试集回归。"),
        ("能力没有被调用", "确认能力已绑定且计数正确；检查能力描述和触发条件；使用更明确的测试问题。"),
        ("长对话逐渐跑题", "降低携带历史轮数或启用历史压缩；在提示词中声明优先级和关键上下文规则。"),
        ("回复被截断", "检查最大回复长度，并让提示词要求分段或先输出摘要；同时关注模型自身限制。"),
        ("发布后用户不可见", "检查“是否发布”状态、用户权限和访问入口；使用普通用户账号复测。"),
        ("外部接入失败", "核对嵌入方式对应的地址、初始化参数、域名策略和认证；不要把访问凭据写入前端或共享文档。"),
    ]
    for question, answer in qa:
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(question)
        set_run_font(r, size=12, bold=True, color=DARK_BLUE)
        doc.add_paragraph(answer)

    doc.add_heading("维护节奏", level=2)
    add_bullet(doc, "模型、提示词或能力变更后，重新执行核心场景回归测试。")
    add_bullet(doc, "定期查看对话记录、用户反馈和点赞点踩数据，识别高频失败场景。")
    add_bullet(doc, "清理不再使用的变量、能力与接入凭据，保持最小权限。")
    add_bullet(doc, "重大变更前保留配置说明和回滚方案，避免直接覆盖稳定版本。")

    add_callout(doc, "页面地址", "http://pro.aiflowy.tech/#/ai/bots/setting/437110544081838080")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
